"""TUBITAK ARDEB Proje Sonuc Raporu bicim kurallarini uygulayan yardimcilar.

Kurallar (psr_kurallari.doc):
- Arial 11 punto, 1,5 satir araligi, tek kolon
- Ust bosluk 3 cm; sag, sol, alt 2,5 cm
- Ana baslik: ortali, BUYUK HARF, koyu, numaradan sonra nokta
- Alt baslik: sola dayali, koyu, iki rakam arasina nokta
- Alt alt baslik: yalnizca ilk harf buyuk
- Tablo basligi tablonun ustunde, sekil basligi seklin altinda; sola dayali,
  en az 10 punto, 1 satir araligi
- Sayfa numarasi ortali ve altta; kapakta yok, on kisimlar romen, ana metin 1'den
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT = "Arial"
BODY_PT = Pt(11)
CAPTION_PT = Pt(10)
BODY_LINE_SPACING = 1.5

# "Ana baslik, alt basliklar ya da paragraflar arasinda 2 satir araligi."
HEADING_SPACE_BEFORE = Pt(24)
HEADING_SPACE_AFTER = Pt(12)


# --------------------------------------------------------------------------- #
# Alan kodlari (Word'un kendi ureteceği icindekiler ve listeler icin)
# --------------------------------------------------------------------------- #
def _field_run(paragraph, instruction: str, *, hidden: bool = False):
    """Word alan kodu ekler. TOC ve TC alanlari bu yolla yazilir."""
    run = paragraph.add_run()
    if hidden:
        rpr = run._element.get_or_add_rPr()
        vanish = OxmlElement("w:vanish")
        rpr.append(vanish)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    element = run._element
    element.append(begin)
    element.append(instr)
    if not hidden:
        element.append(separate)
    element.append(end)
    return run


def _page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    run.font.name = FONT
    run.font.size = BODY_PT
    for tag, attr, text in (
        ("w:fldChar", "begin", None),
        ("w:instrText", None, "PAGE  \\* MERGEFORMAT"),
        ("w:fldChar", "separate", None),
        ("w:fldChar", "end", None),
    ):
        element = OxmlElement(tag)
        if attr:
            element.set(qn("w:fldCharType"), attr)
        if text:
            element.set(qn("xml:space"), "preserve")
            element.text = text
        run._element.append(element)


def _set_page_numbering(section, fmt: str | None, start: int | None) -> None:
    """Bolumun sayfa numarasi bicimini ve baslangicini belirler.

    python-docx yeni bolum eklerken onceki bolumun sectPr'ini kopyalar; bu
    nedecen mevcut pgNumType once kaldirilmalidir, aksi halde Word ilkini
    okur ve ayar etkisiz kalir. Ayrica OOXML sema sirasi geregi pgNumType,
    w:cols ogesinden once gelmelidir.
    """
    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:pgNumType")):
        sect_pr.remove(existing)

    pg = OxmlElement("w:pgNumType")
    if fmt:
        pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))

    anchor = sect_pr.find(qn("w:cols"))
    if anchor is not None:
        anchor.addprevious(pg)
    else:
        sect_pr.append(pg)


# --------------------------------------------------------------------------- #
# Belge iskeleti
# --------------------------------------------------------------------------- #
def new_document() -> Document:
    document = Document()
    _style_normal(document)
    _configure_section(document.sections[0])
    return document


def _style_normal(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_PT
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    paragraph_format = normal.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = BODY_LINE_SPACING
    paragraph_format.space_after = Pt(6)

    # Baslik stilleri de Arial olmali; Word varsayilani mavi Calibri'dir.
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.font.name = FONT
        style.font.bold = True
        style.font.color.rgb = None
        style.font.size = BODY_PT
        style.paragraph_format.space_before = HEADING_SPACE_BEFORE
        style.paragraph_format.space_after = HEADING_SPACE_AFTER
        style.paragraph_format.line_spacing = BODY_LINE_SPACING
        style.paragraph_format.keep_with_next = True


def _configure_section(section) -> None:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_section(document: Document, *, fmt: str | None, start: int | None,
                numbered: bool = True):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(section)
    section.footer.is_linked_to_previous = False
    _set_page_numbering(section, fmt, start)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if numbered:
        _page_number_field(footer)
    return section


# --------------------------------------------------------------------------- #
# Icerik ogeleri
# --------------------------------------------------------------------------- #
def turkish_upper(text: str) -> str:
    """Turkce duyarli buyuk harfe cevirme.

    Python'un str.upper() metodu noktali i harfini I yapar; Turkce'de dogrusu
    İ'dir. Duzeltilmezse "GIRIS", "ICINDEKILER" gibi hatali basliklar olusur.
    """
    return text.replace("i", "İ").replace("ı", "I").upper()


def heading1(document: Document, text: str):
    """Ana konu basligi: ortali, BUYUK HARF, koyu."""
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(turkish_upper(text))
    run.font.name = FONT
    run.font.size = BODY_PT
    run.bold = True
    return paragraph


def heading2(document: Document, text: str):
    """Alt konu basligi: sola dayali, koyu."""
    paragraph = document.add_paragraph(style="Heading 2")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_PT
    run.bold = True
    return paragraph


def heading3(document: Document, text: str):
    """Alt alt konu basligi: ilk harf buyuk, digerleri kucuk."""
    paragraph = document.add_paragraph(style="Heading 3")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_PT
    run.bold = True
    return paragraph


def body(document: Document, text: str, *, justify: bool = True):
    paragraph = document.add_paragraph()
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_PT
    return paragraph


def bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_PT
    return paragraph


def _caption(document: Document, text: str, list_flag: str):
    """Sola dayali, 10 punto, 1 satir araligi baslik + gizli TC alani.

    TC alani, Word'un tablo/sekil listesini dogru sayfa numaralariyla
    uretebilmesi icindir; ekranda gorunmez.
    """
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = CAPTION_PT
    escaped = text.replace('"', "'")
    _field_run(paragraph, f' TC "{escaped}" \\f {list_flag} ', hidden=True)
    return paragraph


def table_caption(document: Document, text: str):
    """Tablo basligi tablonun USTUNDE yer alir."""
    return _caption(document, text, "t")


def figure_caption(document: Document, text: str):
    """Sekil basligi seklin ALTINDA yer alir."""
    return _caption(document, text, "g")


def data_table(document: Document, header: list[str], rows: list[list[str]],
               *, widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = True

    for index, label in enumerate(header):
        cell = table.rows[0].cells[index]
        _fill_cell(cell, label, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            _fill_cell(cells[index], value)

    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    return table


def _fill_cell(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.space_before = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = CAPTION_PT
    run.bold = bold


def figure(document: Document, image_path: Path, width_cm: float = 15.0):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return paragraph


def toc(document: Document, instruction: str):
    paragraph = document.add_paragraph()
    _field_run(paragraph, instruction)
    return paragraph


def page_break(document: Document) -> None:
    document.add_page_break()
