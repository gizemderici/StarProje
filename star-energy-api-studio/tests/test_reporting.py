from __future__ import annotations

import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from reporting import docx_kit
from reporting import sonuc_raporu

ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "data/rapor/Proje_Sonuc_Raporu.docx"


class TurkishUpperTests(unittest.TestCase):
    """Python'un str.upper() metodu Turkce noktali i harfini bozar.

    Duzeltilmezse rapor basliklari "GIRIS" ve "ICINDEKILER" olarak cikar.
    """

    def test_dotted_i_becomes_capital_dotted_i(self) -> None:
        self.assertEqual(docx_kit.turkish_upper("Giriş"), "GİRİŞ")
        self.assertEqual(docx_kit.turkish_upper("İçindekiler"), "İÇİNDEKİLER")
        self.assertEqual(docx_kit.turkish_upper("Sonuç ve Öneriler"), "SONUÇ VE ÖNERİLER")

    def test_dotless_i_becomes_capital_dotless_i(self) -> None:
        self.assertEqual(docx_kit.turkish_upper("Tartışma"), "TARTIŞMA")
        self.assertEqual(docx_kit.turkish_upper("Bulgular"), "BULGULAR")


@unittest.skipUnless(REPORT.exists(), "rapor henuz uretilmemis")
class ReportFormatTests(unittest.TestCase):
    """TUBITAK ARDEB bicim kurallarina uygunluk."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.document = Document(REPORT)

    def test_page_setup_matches_the_rules(self) -> None:
        for section in self.document.sections:
            self.assertAlmostEqual(section.top_margin.cm, 3.0, places=1)
            self.assertAlmostEqual(section.bottom_margin.cm, 2.5, places=1)
            self.assertAlmostEqual(section.left_margin.cm, 2.5, places=1)
            self.assertAlmostEqual(section.right_margin.cm, 2.5, places=1)

    def test_body_font_is_arial_eleven_at_one_and_a_half_spacing(self) -> None:
        normal = self.document.styles["Normal"]
        self.assertEqual(normal.font.name, "Arial")
        self.assertEqual(normal.font.size.pt, 11)
        self.assertEqual(normal.paragraph_format.line_spacing, 1.5)

    def test_cover_has_no_page_number_and_front_matter_is_roman(self) -> None:
        # Kapakta sayfa numarasi olmamali, on kisimlar romen, ana metin 1'den.
        formats = []
        for section in self.document.sections:
            entries = section._sectPr.findall(qn("w:pgNumType"))
            # Kopyalanan sectPr yuzunden birden fazla olusabilir; olmamali.
            self.assertLessEqual(len(entries), 1)
            formats.append(entries[0].get(qn("w:fmt")) if entries else None)
        self.assertEqual(formats, [None, "lowerRoman", "decimal"])

    def test_captions_are_at_least_ten_point_single_spaced(self) -> None:
        # Alanlar guncellendikten sonra Tablo/Sekil Listesi girdileri de
        # "Tablo ..." ile baslar; onlar "toc" stilindedir ve 10 punto kurali
        # baslik paragraflari icindir, liste girdileri icin degil.
        captions = [
            paragraph
            for paragraph in self.document.paragraphs
            if paragraph.text.startswith(("Tablo ", "Şekil "))
            and not paragraph.style.name.lower().startswith("toc")
        ]
        self.assertGreater(len(captions), 10)
        for paragraph in captions:
            self.assertGreaterEqual(paragraph.runs[0].font.size.pt, 10)
            self.assertEqual(paragraph.paragraph_format.line_spacing, 1.0)

    def test_main_headings_are_centred_and_correctly_capitalised(self) -> None:
        headings = [
            paragraph.text
            for paragraph in self.document.paragraphs
            if paragraph.style.name == "Heading 1"
        ]
        for required in ("1. GİRİŞ", "4. BULGULAR", "6. SONUÇ VE ÖNERİLER",
                         "KAYNAKLAR", "İÇİNDEKİLER"):
            self.assertIn(required, headings)

    def test_required_sections_follow_the_template_order(self) -> None:
        headings = [
            paragraph.text
            for paragraph in self.document.paragraphs
            if paragraph.style.name == "Heading 1"
        ]
        expected = [
            "ÖNSÖZ", "İÇİNDEKİLER", "TABLO LİSTESİ", "ŞEKİL LİSTESİ",
            "ÖZET", "ABSTRACT", "1. GİRİŞ", "2. LİTERATÜR ÖZETİ",
            "3. GEREÇ VE YÖNTEM", "4. BULGULAR", "5. TARTIŞMA",
            "6. SONUÇ VE ÖNERİLER", "KAYNAKLAR", "EKLER",
        ]
        self.assertEqual(headings, expected)

    def test_abstracts_stay_within_the_word_limit(self) -> None:
        texts = [paragraph.text for paragraph in self.document.paragraphs]
        for title in ("ÖZET", "ABSTRACT"):
            start = texts.index(title)
            words: list[str] = []
            for text in texts[start + 1:]:
                if text.startswith(("Anahtar Kelimeler", "Keywords")):
                    break
                words.extend(text.split())
            self.assertGreaterEqual(len(words), 100, f"{title} cok kisa")
            self.assertLessEqual(len(words), 300, f"{title} cok uzun")

    def test_figures_are_embedded(self) -> None:
        # 4 analiz grafigi + 1 mimari sema + 7 arayuz ekran goruntusu
        self.assertEqual(len(self.document.inline_shapes), 12)

    def test_interface_screenshots_are_included(self) -> None:
        captions = [
            paragraph.text
            for paragraph in self.document.paragraphs
            if paragraph.text.startswith("Şekil ")
        ]
        for required in ("Şekil 3.1.", "Şekil 3.2.", "Şekil 4.5.", "Şekil 4.8."):
            self.assertTrue(
                any(caption.startswith(required) for caption in captions),
                f"{required} eksik",
            )


class ReportContentTests(unittest.TestCase):
    """Rapordaki sayilar uretilmis analiz dosyalariyla ayni olmalidir."""

    def test_facts_are_read_from_generated_reports(self) -> None:
        facts = sonuc_raporu.collect_facts()
        energy = facts["scores"]["site_energy_gj"]
        self.assertTrue(energy["meets_target"])
        self.assertTrue(facts["summary"]["within_tolerance"])
        # TOPSIS uzlasi cozumu daima dogrulanan noktalar arasinda olmalidir.
        self.assertIn("TOPSIS", facts["topsis"]["reason"])

    def test_thesis_claim_about_insulation_still_holds(self) -> None:
        # Raporun ana savi: yalitim kalinligi son sirada, chiller COP ilk sirada.
        facts = sonuc_raporu.collect_facts()
        ordered = sorted(
            facts["sensitivity"], key=lambda item: item["total"], reverse=True
        )
        self.assertEqual(ordered[0]["key"], "chiller_cop")
        self.assertEqual(ordered[-1]["key"], "eps_thickness_cm")


if __name__ == "__main__":
    unittest.main()
